import os
import json
import io
from pydantic import BaseModel, Field, FieldValidationInfo, field_validator
import docx
from docx.shared import Inches, Pt

# ==========================================
# ARCHITECTURAL DESIGN: Pydantic Schema
# ==========================================
class ReporteSlots(BaseModel):
    # Slots Determinísticos (Mapping directo sin LLM)
    id_iniciativa: int
    titulo: str
    facultad: str
    carrera: str
    
    # Slots Generativos (A ser procesados por LLM con control de longitud)
    resumen_ejecutivo: str = Field(..., min_length=50, max_length=500)
    conclusiones_vcm: str = Field(..., min_length=50)
    
    # Métrica de calidad RAGAS pre-guardado
    ragas_faithfulness: float = Field(..., ge=0.0, le=1.0)

    @field_validator('ragas_faithfulness')
    @classmethod
    def verificar_umbral_calidad(cls, v: float) -> float:
        UMBRAL_MINIMO = 0.75 # Definido en Specs de G3
        if v < UMBRAL_MINIMO:
            raise ValueError(f"Calidad insuficiente ({v}). Requiere revisión humana obligatoria.")
        return v

# ==========================================
# MOTOR DOCX: Generación de la Plantilla Base
# ==========================================
def construir_documento_docx(datos: ReporteSlots) -> io.BytesIO:
    doc = docx.Document()
    
    # Título Estándar Institucional UTEM
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = 1 # Center
    run_t = p_titulo.add_run("INFORME DE GESTIÓN Y EVIDENCIAS VcM\nUNIVERSIDAD TECNOLÓGICA METROPOLITANA")
    run_t.bold = True
    run_t.font.size = Pt(12)
    
    doc.add_paragraph(f"ID Iniciativa: {datos.id_iniciativa}").alignment = 2 # Right
    doc.add_paragraph("-" * 60)
    
    # 1. Datos Generales (Determinísticos)
    h1 = doc.add_paragraph()
    h1.add_run("1. ANTECEDENTES DE LA INICIATIVA (Mapping Automático)").bold = True
    doc.add_paragraph(f"• Título: {datos.titulo}")
    doc.add_paragraph(f"• Unidad: {datos.facultad} / {datos.carrera}")
    
    # 2. Secciones del LLM (Generativos)
    h2 = doc.add_paragraph()
    h2.add_run("\n2. RESUMEN NARRATIVO (Generación Controlada LLM)").bold = True
    doc.add_paragraph(datos.resumen_ejecutivo)
    
    h3 = doc.add_paragraph()
    h3.add_run("\n3. CONCLUSIONES E IMPACTO").bold = True
    doc.add_paragraph(datos.conclusiones_vcm)
    
    # Footer de control de calidad
    doc.add_paragraph("-" * 60)
    footer = doc.add_paragraph()
    footer.add_run(f"Métrica RAGAS Faithfulness: {datos.ragas_faithfulness} (Umbral de confianza: OK)").italic = True
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

if __name__ == "__main__":
    print("Prueba del Pipeline de Datos del Grupo 3...")
    # Datos de simulación post-LLM con Score de RAGAS simulado
    mock_llm_output = {
        "id_iniciativa": 2606,
        "titulo": "Optimización de Procesos de Datos para Analistas VcM UTEM",
        "facultad": "Facultad de Ingeniería",
        "carrera": "Ingeniería Civil en Ciencia de Datos",
        "resumen_ejecutivo": "Este documento técnico detalla la implementación del pipeline automatizado del Grupo 3, diseñado para reducir la carga de trabajo operativo de los analistas de Vinculación con el Medio de la UTEM mediante el uso de inteligencia artificial adaptativa.",
        "conclusiones_vcm": "La automatización respeta el criterio profesional del analista mediante una UI de Streamlit, actuando como un borrador de alta calidad y reduciendo drásticamente los tiempos de procesamiento institucional.",
        "ragas_faithfulness": 0.89  # Pasa el umbral de 0.75
    }
    
    validado = ReporteSlots(**mock_llm_output)
    archivo = construir_documento_docx(validado)
    print("✅ Pipeline ejecutado: Objeto Pydantic validado y binario .docx estructurado exitosamente.")